"""
Generate Laporan PKL Word Document (.docx) - FTIK Universitas Teknokrat Indonesia
Diperbarui dengan data resmi Tim & Profil CV Newus Teknologi (https://newus.id/team)
serta data mahasiswa yang diperbaiki oleh user:
1. FERY DWI RAMADHI (23312086)
2. FATHUR RAMANTHA (23312105)
3. I PUTU PANDU WIRANATA (23312088)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def sfmt(paragraph, space_before=0, space_after=0, line_spacing=1.5, 
         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None,
         keep_with_next=False, left_indent=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    if keep_with_next:
        pf.keep_with_next = True

def ar(paragraph, text, bold=False, italic=False, size=12, underline=False, color=None):
    r = paragraph.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r

def set_page_number_roman(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE  \\* ROMANLOW </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

def set_page_number_arabic(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    footer = section.footer
    footer.is_linked_to_previous = False

def add_gambar(doc, placeholder_text, caption):
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_before=12, space_after=6)
    ar(p, f"[Gambar untuk {placeholder_text}]", italic=True, size=11)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, caption, size=12)

def add_tabel_caption(doc, caption):
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_before=6, keep_with_next=True)
    ar(p, caption, size=12)

def add_blue_separator(doc, name, npm):
    doc.add_page_break()
    for _ in range(8):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, space_after=0)
    ar(p, "BAB III", bold=True, size=14, color=(0, 32, 96))
    
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, space_after=12)
    ar(p, "PELAKSANAAN PRAKTIK KERJA LAPANGAN", bold=True, size=14, color=(0, 32, 96))
    
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=0)
    ar(p, name, bold=True, size=16, color=(0, 32, 96))
    
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=0)
    ar(p, npm, bold=True, size=14, color=(0, 32, 96))
    
    doc.add_page_break()

def add_bab_title(doc, bab_num, bab_text):
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_before=28)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0, keep_with_next=True)
    ar(p, f"BAB {bab_num}", bold=True, size=12)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12, keep_with_next=True)
    ar(p, bab_text, bold=True, size=12)

def add_subbab(doc, number, title):
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_before=6, space_after=6, keep_with_next=True)
    ar(p, f"{number} {title}", bold=True, size=12)

def add_body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    sfmt(p, line_spacing=1.5, first_line_indent=1.27 if first_indent else None)
    ar(p, text, size=12)
    return p

def add_body_mixed(doc, first_indent=True):
    p = doc.add_paragraph()
    sfmt(p, line_spacing=1.5, first_line_indent=1.27 if first_indent else None)
    return p

def add_numbered_item(doc, number, text, indent=0):
    p = doc.add_paragraph()
    sfmt(p, line_spacing=1.5, space_after=0, left_indent=indent if indent else None)
    ar(p, f"{number}) {text}", size=12)

def create_laporan():
    doc = Document()

    # GLOBAL STYLES
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # SECTION 1: Bagian Awal (Romawi Kecil di Bawah Tengah)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    set_page_number_roman(section)

    names_data = [
        ("1.", "FERY DWI RAMADHI", "(23312086)"),
        ("2.", "FATHUR RAMANTHA", "(23312105)"),
        ("3.", "I PUTU PANDU WIRANATA", "(23312088)"),
    ]

    # =========================================================================
    # 1. HALAMAN SAMPUL DAN JUDUL (Lampiran 1)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_before=24)

    for line in ["LAPORAN PRAKTIK KERJA LAPANGAN (PKL)", "DI CV NEWUS TEKNOLOGI", "BANDAR LAMPUNG"]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        ar(p, line, bold=True, size=14)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    ar(p, "Untuk memenuhi persyaratan mendapatkan nilai Praktik Kerja Lapangan", size=12)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
    ar(p, "Disusun oleh:", bold=True, size=12)

    for num, name, npm in names_data:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        ar(p, f"{num} {name} {npm}", bold=True, size=12)

    for _ in range(2):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    ar(p, "[Gambar untuk Logo Universitas Teknokrat Indonesia]", italic=True, size=11)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    ar(p, "(Tinggi gambar 5 cm, Lebar gambar 5,5 cm)", italic=True, size=10)

    for _ in range(2):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    for line in ["PROGRAM STUDI S1 INFORMATIKA", "FAKULTAS TEKNIK DAN ILMU KOMPUTER",
                 "UNIVERSITAS TEKNOKRAT INDONESIA", "BANDAR LAMPUNG", "2026"]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        ar(p, line, bold=True, size=14)

    doc.add_page_break()

    # =========================================================================
    # 2. LEMBAR PERSETUJUAN (Lampiran 2)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
    ar(p, "LEMBAR PERSETUJUAN LAPORAN", bold=True, size=14)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "PRAKTIK KERJA LAPANGAN", bold=True, size=14)

    info = [
        ("Nama", ":\t1.  Fery Dwi Ramadhi (23312086)"),
        ("", "\t2.  Fathur Ramantha (23312105)"),
        ("", "\t3.  I Putu Pandu Wiranata (23312088)"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, space_after=0)
        if label:
            ar(p, label, size=12)
        ar(p, value, size=12)

    for label, value in [
        ("Program Studi", ":\tS1 Informatika"), 
        ("Instansi/perusahaan", ":\tCV Newus Teknologi"),
        ("Alamat Instansi/perusahaan", ":\tJl. Salim Batubara No.118, Kupang Teba,"),
        ("", "\t\tKec. Teluk Betung Utara, Bandar Lampung")
    ]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, space_after=0)
        if label:
            ar(p, label, size=12)
            ar(p, f"\t{value}", size=12)
        else:
            ar(p, value, size=12)

    for _ in range(2):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c = tbl.rows[0].cells[0].paragraphs[0]
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(c, "Pembimbing,", size=12)
    
    c1 = tbl.rows[1].cells[0].paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(c1, "Pembimbing laporan PKL\nFakultas Teknik dan Ilmu Komputer,", size=12)
    c2 = tbl.rows[1].cells[1].paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(c2, "Pembimbing lapangan\nInstansi/Perusahaan PKL,", size=12)

    for i in range(2):
        c = tbl.rows[2].cells[i].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, "\n\n\n\n", size=12)
    
    for i in range(2):
        c = tbl.rows[3].cells[i].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, "(.......................................................)", size=12)
        px = tbl.rows[3].cells[i].add_paragraph()
        px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(px, "NIK", size=12)

    for row in tbl.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    for _ in range(2):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    for text in ["Menyetujui,", "Program Studi Informatika", "Ketua,"]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        ar(p, text, size=12)

    for _ in range(3):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
    ar(p, "Dr. Heni Sulistiani, S.Kom., M.Kom.", size=12)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    ar(p, "NIK 022 13 02 11", size=12)

    doc.add_page_break()

    # =========================================================================
    # 3. LEMBAR PENGESAHAN (Lampiran 3)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "LEMBAR PENGESAHAN", bold=True, size=14)

    for line in ["LAPORAN PRAKTIK KERJA LAPANGAN (PKL)", "DI CV NEWUS TEKNOLOGI", "BANDAR LAMPUNG"]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        ar(p, line, bold=True, size=12)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_before=6)
    ar(p, "Dipersiapkan dan disusun oleh:", size=12)

    for num, name, npm in names_data:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
        ar(p, f"{num}   {name}  {npm}", bold=True, size=12)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_before=6)
    ar(p, "Telah dipertahankan di depan Dewan Penguji", size=12)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    ar(p, "Pada tanggal 6 September 2026", size=12)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_before=6)
    ar(p, "Dewan Penguji", bold=True, size=12)

    tbl2 = doc.add_table(rows=3, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(["Pembimbing,", "Penguji,"]):
        c = tbl2.rows[0].cells[idx].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, label, size=12)
    for i in range(2):
        c = tbl2.rows[1].cells[i].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, "\n\n\n\n", size=12)
    for i in range(2):
        c = tbl2.rows[2].cells[i].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, "(.......................................................)", size=12)
        px = tbl2.rows[2].cells[i].add_paragraph()
        px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(px, "NIK", size=12)
    for row in tbl2.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_before=6)
    ar(p, "Laporan ini telah diterima sebagai salah satu persyaratan", size=12)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
    ar(p, "untuk memperoleh nilai Praktik Kerja Lapangan", size=12)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    ar(p, "Tanggal 6 September 2026", size=12)

    tbl3 = doc.add_table(rows=3, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, lines in enumerate([
        ("Fakultas Teknik dan Ilmu Komputer", "Dekan,"),
        ("Program Studi Informatika", "Ketua,")
    ]):
        c = tbl3.rows[0].cells[idx].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, lines[0], size=12)
        px = tbl3.rows[0].cells[idx].add_paragraph()
        px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(px, lines[1], size=12)

    for i in range(2):
        c = tbl3.rows[1].cells[i].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c, "\n\n\n\n", size=12)

    c = tbl3.rows[2].cells[0].paragraphs[0]
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(c, "Dr.Si. Dedi Darwis, M.Kom., CDSP.", size=12)
    px = tbl3.rows[2].cells[0].add_paragraph()
    px.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(px, "NIK 023 05 00 09", size=12)

    c = tbl3.rows[2].cells[1].paragraphs[0]
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(c, "Dr. Heni Sulistiani, M.Kom.", size=12)
    px = tbl3.rows[2].cells[1].add_paragraph()
    px.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(px, "NIK 022 13 02 11", size=12)

    for row in tbl3.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    doc.add_page_break()

    # =========================================================================
    # 4. KATA PENGANTAR (Lampiran 4)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "KATA PENGANTAR", bold=True, size=14)

    add_body(doc, "Puji syukur penulis panjatkan kepada Allah SWT, karena atas berkat dan rahmat-Nya, penulis dapat menyelesaikan Laporan Praktik Kerja Lapangan (PKL) ini.")
    add_body(doc, "Penulisan Laporan PKL ini dilakukan dalam rangka memenuhi salah satu syarat untuk mendapatkan nilai Praktik Kerja Lapangan (PKL) pada Program Studi S1 Informatika Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia. Penulis menyadari bahwa, tanpa bantuan dan bimbingan dari berbagai pihak, sangatlah sulit bagi penulis untuk menyelesaikan laporan PKL ini. Oleh karena itu, penulis mengucapkan terima kasih kepada:")

    thanks = [
        "Dr. H.M. Nasrullah Yusuf, S.E., M.B.A. selaku Rektor Universitas Teknokrat Indonesia.",
        "Dr.Si. Dedi Darwis, M.Kom., CDSP., selaku Dekan Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia.",
        "Dr. Heni Sulistiani, M.Kom., selaku Ketua Program Studi S1 Informatika Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia.",
        "Yusra Fernando, M.Kom., selaku Koordinator PKL Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia.",
        "Dosen Pembimbing yang telah meluangkan waktu untuk membimbing penulis menyelesaikan Laporan PKL ini.",
        "Dosen Penguji yang telah memberikan evaluasi, masukan, dan saran yang membangun bagi kesempurnaan laporan PKL ini.",
        "Ir. Antoni, S.T., IPM selaku Founder & CEO CV Newus Teknologi beserta seluruh tim profesional yang telah banyak membantu, membimbing, dan memberikan kesempatan berharga selama melaksanakan PKL.",
        "Orang tua dan keluarga tercinta yang senantiasa memberikan doa, motivasi, dan dukungan moral maupun materiil.",
        "Seluruh rekan mahasiswa Program Studi S1 Informatika Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia.",
    ]
    for i, t in enumerate(thanks, 1):
        add_numbered_item(doc, i, t)

    add_body(doc, "Akhir kata, penulis berharap semoga Allah SWT berkenan membalas segala kebaikan semua pihak yang telah membantu dan semoga Laporan PKL ini membawa manfaat bagi pengembangan ilmu pengetahuan dan teknologi.")

    for _ in range(2):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0, space_after=0)
    ar(p, "Bandar Lampung, 6 September 2026", size=12)
    for _ in range(3):
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0)
    ar(p, "Penulis", size=12)

    doc.add_page_break()

    # =========================================================================
    # 5. DAFTAR ISI (Lampiran 5)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "DAFTAR ISI", bold=True, size=14)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0, space_after=6)
    ar(p, "Hal", size=12)

    daftar_isi = [
        ("LEMBAR JUDUL", "i", False, 0),
        ("LEMBAR PERSETUJUAN", "ii", False, 0),
        ("LEMBAR PENGESAHAN", "iii", False, 0),
        ("KATA PENGANTAR", "iv", False, 0),
        ("DAFTAR ISI", "v", False, 0),
        ("DAFTAR TABEL", "vi", False, 0),
        ("DAFTAR GAMBAR", "vii", False, 0),
        ("DAFTAR LAMPIRAN", "viii", False, 0),
        ("RINGKASAN PELAKSANAAN PKL", "ix", False, 0),
        ("BAB I PENDAHULUAN", "1", True, 0),
        ("1.1 Latar Belakang", "1", False, 0.63),
        ("1.2 Tujuan PKL", "3", False, 0.63),
        ("1.3 Kegunaan PKL", "4", False, 0.63),
        ("1.4 Tempat PKL", "6", False, 0.63),
        ("1.5 Jadwal Pelaksanaan PKL", "7", False, 0.63),
        ("BAB II TINJAUAN UMUM TEMPAT PKL", "9", True, 0),
        ("2.1 Sejarah Perusahaan", "9", False, 0.63),
        ("2.2 Struktur Organisasi", "11", False, 0.63),
        ("2.3 Kegiatan Umum Perusahaan", "13", False, 0.63),
        ("BAB III PELAKSANAAN PRAKTIK KERJA LAPANGAN", "15", True, 0),
        ("3.1 FERY DWI RAMADHI", "15", True, 0.63),
        ("3.1.1 Bidang Kerja", "15", False, 1.27),
        ("3.1.2 Pelaksanaan Kerja", "16", False, 1.27),
        ("3.1.3 Kendala yang Dihadapi", "25", False, 1.27),
        ("3.1.4 Cara Mengatasi Kendala", "26", False, 1.27),
        ("3.2 FATHUR RAMANTHA", "29", True, 0.63),
        ("3.2.1 Bidang Kerja", "29", False, 1.27),
        ("3.2.2 Pelaksanaan Kerja", "30", False, 1.27),
        ("3.2.3 Kendala Yang Dihadapi", "37", False, 1.27),
        ("3.2.4 Cara Mengatasi Kendala", "38", False, 1.27),
        ("3.3 I PUTU PANDU WIRANATA", "41", True, 0.63),
        ("3.3.1 Bidang Kerja", "41", False, 1.27),
        ("3.3.2 Pelaksanaan Kerja", "42", False, 1.27),
        ("3.3.3 Kendala Yang Dihadapi", "49", False, 1.27),
        ("3.3.4 Cara Mengatasi Kendala", "50", False, 1.27),
        ("BAB IV PENUTUP", "53", True, 0),
        ("4.1 Simpulan", "53", False, 0.63),
        ("4.2 Saran", "55", False, 0.63),
        ("DAFTAR PUSTAKA", "57", True, 0),
        ("LAMPIRAN", "58", True, 0),
    ]
    for item, page, is_bold, indent in daftar_isi:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=0, left_indent=indent if indent else None)
        ar(p, item, bold=is_bold, size=12)
        dots = max(3, 60 - len(item) - int(indent*3))
        ar(p, f"  {'.' * dots}  {page}", size=12)

    doc.add_page_break()

    # =========================================================================
    # 6. DAFTAR TABEL (Lampiran 5)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "DAFTAR TABEL", bold=True, size=14)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0, space_after=6)
    ar(p, "Hal", size=12)

    for t_no, t_judul, t_hal in [
        ("Tabel 1.1", "Tahapan Kegiatan Praktik Kerja Lapangan (PKL)", "8"),
        ("Tabel 2.1", "Daftar Susunan Tim Inti CV Newus Teknologi", "12"),
        ("Tabel 3.1", "Daftar Endpoint RESTful API Laravel 11", "19"),
        ("Tabel 3.2", "Struktur Tabel Basis Data MySQL", "17"),
        ("Tabel 3.3", "Daftar Komponen Antarmuka Next.js", "32"),
        ("Tabel 3.4", "Hasil Pengujian Otomatis (E2E Automated Bot Testing)", "47"),
    ]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=0)
        item_text = f"{t_no} {t_judul}"
        ar(p, item_text, size=12)
        dots = max(3, 62 - len(item_text))
        ar(p, f"  {'.' * dots}  {t_hal}", size=12)

    doc.add_page_break()

    # =========================================================================
    # 7. DAFTAR GAMBAR (Lampiran 5)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "DAFTAR GAMBAR", bold=True, size=14)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0, space_after=6)
    ar(p, "Hal", size=12)

    for g_no, g_judul, g_hal in [
        ("Gambar 1.1", "Tempat Praktik Kerja Lapangan (PKL)", "7"),
        ("Gambar 2.1", "Logo CV Newus Teknologi", "10"),
        ("Gambar 2.2", "Struktur Organisasi CV Newus Teknologi", "12"),
        ("Gambar 3.1", "Arsitektur Sistem LMS (Frontend dan Backend)", "15"),
        ("Gambar 3.2", "Entity Relationship Diagram (ERD) Basis Data", "16"),
        ("Gambar 3.3", "Diagram Use Case Terpadu Sistem LMS", "43"),
        ("Gambar 3.4", "Diagram Alir (Flowchart) Sistem Terpadu", "44"),
        ("Gambar 3.5", "Tampilan Halaman Login Sistem LMS", "21"),
        ("Gambar 3.6", "Tampilan Dashboard Administrator", "31"),
        ("Gambar 3.7", "Tampilan Dashboard Guru Pengajar", "33"),
        ("Gambar 3.8", "Tampilan Dashboard Siswa", "34"),
        ("Gambar 3.9", "Tampilan Halaman Manajemen Kelas", "35"),
        ("Gambar 3.10", "Tampilan Halaman Presensi Siswa", "36"),
        ("Gambar 3.11", "Tampilan Halaman Penugasan LKPD", "36"),
        ("Gambar 3.12", "Tampilan Halaman Rekapitulasi Nilai Rapor", "45"),
        ("Gambar 3.13", "Tampilan Lonceng Notifikasi Real-time", "45"),
        ("Gambar 3.14", "Tampilan Fitur Ekspor Laporan Excel", "46"),
        ("Gambar 3.15", "Tampilan Halaman Manajemen Modul Bahan Ajar", "46"),
        ("Gambar 3.16", "Tampilan Halaman Input Nilai UTS/UAS", "47"),
        ("Gambar 3.17", "Tampilan Halaman Bulk Import Spreadsheet", "44"),
        ("Gambar 3.18", "Dokumentasi Kegiatan PKL di CV Newus Teknologi", "48"),
    ]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=0)
        item_text = f"{g_no} {g_judul}"
        ar(p, item_text, size=12)
        dots = max(3, 62 - len(item_text))
        ar(p, f"  {'.' * dots}  {g_hal}", size=12)

    doc.add_page_break()

    # =========================================================================
    # 8. DAFTAR LAMPIRAN (Lampiran 5)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "DAFTAR LAMPIRAN", bold=True, size=14)

    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0, space_after=6)
    ar(p, "Hal", size=12)

    for l_no, l_judul, l_hal in [
        ("Lampiran 1", "Formulir Penilaian Fery Dwi Ramadhi", "58"),
        ("Lampiran 2", "Catatan Harian Fery Dwi Ramadhi", "59"),
        ("Lampiran 3", "Formulir Penilaian Fathur Ramantha", "60"),
        ("Lampiran 4", "Catatan Harian Fathur Ramantha", "61"),
        ("Lampiran 5", "Formulir Penilaian I Putu Pandu Wiranata", "62"),
        ("Lampiran 6", "Catatan Harian I Putu Pandu Wiranata", "63"),
    ]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, space_after=0)
        item_text = f"{l_no} {l_judul}"
        ar(p, item_text, size=12)
        dots = max(3, 62 - len(item_text))
        ar(p, f"  {'.' * dots}  {l_hal}", size=12)

    doc.add_page_break()

    # =========================================================================
    # 9. RINGKASAN PELAKSANAAN PRAKTIK KERJA LAPANGAN (Lampiran 6)
    # =========================================================================
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=0)
    ar(p, "RINGKASAN PELAKSANAAN", bold=True, size=14)
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
    ar(p, "PRAKTIK KERJA LAPANGAN", bold=True, size=14)

    p = add_body_mixed(doc)
    p.paragraph_format.line_spacing = 1.0
    ar(p, "Praktik Kerja Lapangan (PKL) merupakan sarana mengaktualisasikan diri terhadap berbagai keahlian dan keterampilan baik ", size=12)
    ar(p, "soft skill ", italic=True, size=12)
    ar(p, "maupun ", size=12)
    ar(p, "hard skill ", italic=True, size=12)
    ar(p, "yang telah diperoleh selama masa perkuliahan di Program Studi S1 Informatika Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia. Kegiatan PKL dilaksanakan di CV Newus Teknologi yang berlokasi di Jl. Salim Batubara No.118, Kupang Teba, Kec. Teluk Betung Utara, Kota Bandar Lampung, Lampung selama satu bulan terhitung mulai tanggal 6 Agustus 2026 sampai dengan 6 September 2026. Penulis yang terdiri atas tiga orang mahasiswa (Fery Dwi Ramadhi, Fathur Ramantha, dan I Putu Pandu Wiranata) ditempatkan pada posisi ", size=12)
    ar(p, "Software Engineer ", italic=True, size=12)
    ar(p, "di bawah bimbingan langsung tim profesional Newus Technology dengan tanggung jawab merancang dan membangun platform ", size=12)
    ar(p, "Learning Management System ", italic=True, size=12)
    ar(p, "(LMS) berbasis web sesuai dengan kebutuhan spesifikasi industri.", size=12)

    p = add_body_mixed(doc)
    p.paragraph_format.line_spacing = 1.0
    ar(p, "Pelaksanaan kerja pada proyek LMS dibagi ke dalam beberapa bidang tanggung jawab: Fery Dwi Ramadhi bertanggung jawab atas perancangan arsitektur sistem, basis data MySQL 11 tabel relasional, implementasi lebih dari 50 ", size=12)
    ar(p, "endpoint RESTful API ", italic=True, size=12)
    ar(p, "menggunakan Laravel 11, serta mekanisme sinkronisasi data ", size=12)
    ar(p, "real-time ", italic=True, size=12)
    ar(p, "antar-tab browser; Fathur Ramantha bertanggung jawab atas perancangan antarmuka pengguna (UI/UX) di Figma, implementasi 30 halaman dan 20 komponen ", size=12)
    ar(p, "reusable ", italic=True, size=12)
    ar(p, "Next.js dengan Tailwind CSS untuk tiga peran pengguna (Administrator, Guru Pengajar, dan Peserta Didik); serta I Putu Pandu Wiranata bertanggung jawab atas analisis kebutuhan fungsional (PRD dan diagram UML), manajemen ", size=12)
    ar(p, "seeder ", italic=True, size=12)
    ar(p, "61 akun pengguna, modul ", size=12)
    ar(p, "bulk import spreadsheet, ", italic=True, size=12)
    ar(p, "dan perancangan pengujian otomatis (", size=12)
    ar(p, "End-to-End Automated Bot Testing", italic=True, size=12)
    ar(p, ").", size=12)

    p = add_body_mixed(doc)
    p.paragraph_format.line_spacing = 1.0
    ar(p, "Hasil pengujian otomatis terhadap 53 skenario uji coba lintas-peran menunjukkan tingkat keberhasilan 100% (", size=12)
    ar(p, "all passed", italic=True, size=12)
    ar(p, "). Seluruh fungsionalitas sistem meliputi manajemen kelas, modul materi, penugasan LKPD ", size=12)
    ar(p, "daring, ", italic=True, size=12)
    ar(p, "presensi mandiri berbasis jendela waktu, penilaian ujian UTS/UAS, lonceng notifikasi ", size=12)
    ar(p, "real-time, ", italic=True, size=12)
    ar(p, "dan ekspor rekapitulasi nilai ke format Excel (.xlsx) telah berhasil di-", size=12)
    ar(p, "deploy ", italic=True, size=12)
    ar(p, "pada infrastruktur ", size=12)
    ar(p, "cloud ", italic=True, size=12)
    ar(p, "produksi (Vercel dan Railway). Kegiatan PKL ini memberikan pengalaman berharga mengenai dinamika kerja industri perangkat lunak, memperkuat kompetensi teknis, serta melatih kerja sama tim secara profesional.", size=12)

    p = doc.add_paragraph()
    sfmt(p, line_spacing=1.0, space_before=12)
    ar(p, "Kata Kunci : ", bold=True, size=12)
    ar(p, "PKL, CV Newus Teknologi, Learning Management System, Next.js, Laravel, MySQL.", italic=True, size=12)

    # =========================================================================
    # SECTION 2: Bagian Utama (Angka Arab di Kanan Atas)
    # =========================================================================
    new_section = doc.add_section()
    new_section.page_width = Cm(21)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(3)
    new_section.bottom_margin = Cm(3)
    new_section.left_margin = Cm(4)
    new_section.right_margin = Cm(3)
    set_page_number_arabic(new_section)

    # =========================================================================
    # BAB I PENDAHULUAN
    # =========================================================================
    add_bab_title(doc, "I", "PENDAHULUAN")

    add_subbab(doc, "1.1", "Latar Belakang")

    add_body(doc, "Perkembangan teknologi informasi yang semakin pesat telah mendorong berbagai sektor industri untuk memanfaatkan sistem informasi dalam mendukung aktivitas operasional. Sektor pendidikan sebagai salah satu pilar utama pembangunan nasional tidak terkecuali dalam transformasi digital tersebut. Institusi pendidikan dituntut untuk mampu mengelola data akademik, proses pembelajaran, dan administrasi sekolah secara cepat, akurat, serta terintegrasi guna meningkatkan efisiensi dan efektivitas pengelolaan pendidikan.")

    p = add_body_mixed(doc)
    ar(p, "Salah satu bentuk pemanfaatan teknologi informasi dalam dunia pendidikan adalah melalui pengembangan ", size=12)
    ar(p, "Learning Management System ", italic=True, size=12)
    ar(p, "(LMS). LMS merupakan perangkat lunak yang dirancang untuk membantu pengelolaan kegiatan pembelajaran secara ", size=12)
    ar(p, "daring, ", italic=True, size=12)
    ar(p, "meliputi pengelolaan kelas, pendistribusian bahan ajar, penugasan, penilaian, hingga pemantauan kehadiran peserta didik. Dengan adanya LMS, proses pembelajaran tidak lagi dibatasi oleh ruang dan waktu, sehingga interaksi antara guru pengajar dan peserta didik dapat berlangsung secara lebih fleksibel dan terdokumentasi dengan baik.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Fakultas Teknik dan Ilmu Komputer (FTIK) Universitas Teknokrat Indonesia memiliki tujuan untuk menghasilkan lulusan yang tidak hanya menguasai teori, tetapi juga mampu mengimplementasikan ilmu pengetahuan dan teknologi secara profesional di dunia kerja. Namun demikian, terdapat kesenjangan antara pembelajaran yang diperoleh selama perkuliahan dengan kondisi nyata di lingkungan industri. Dalam dunia kerja, seorang tenaga profesional tidak hanya dituntut memiliki kemampuan teknis, tetapi juga kemampuan menganalisis kebutuhan pengguna, berkomunikasi dengan klien, bekerja sama dalam tim, serta beradaptasi terhadap perubahan kebutuhan sistem yang dinamis. Oleh karena itu, diperlukan suatu kegiatan yang dapat menjembatani mahasiswa dengan dunia kerja profesional, yaitu melalui Praktik Kerja Lapangan (PKL).", size=12)

    p = add_body_mixed(doc)
    ar(p, "CV Newus Teknologi dipilih sebagai tempat pelaksanaan Praktik Kerja Lapangan karena merupakan perusahaan konsultan teknologi informasi dan pengembang perangkat lunak berpengalaman yang berbasis di Bandar Lampung. Perusahaan ini telah berdiri sejak tahun 2021 di bawah kepemimpinan Ir. Antoni, S.T., IPM dan telah menyelesaikan lebih dari 100 solusi teknologi digital untuk sektor pemerintahan, kesehatan, pendidikan, dan korporasi. Lingkungan kerja Newus Technology sangat relevan dengan bidang keilmuan Informatika, khususnya dalam pengembangan sistem informasi berskala besar berbasis arsitektur modern.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Selama melaksanakan PKL, penulis ditugaskan untuk merancang dan membangun ", size=12)
    ar(p, "Learning Management System ", italic=True, size=12)
    ar(p, "(LMS) berbasis web menggunakan ", size=12)
    ar(p, "framework ", italic=True, size=12)
    ar(p, "Next.js pada sisi ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "dan Laravel pada sisi ", size=12)
    ar(p, "backend ", italic=True, size=12)
    ar(p, "yang terintegrasi dengan basis data MySQL.", size=12)

    add_body(doc, "Melalui pelaksanaan PKL di CV Newus Teknologi, penulis memperoleh pengalaman nyata mengenai proses pengembangan sistem informasi yang berorientasi pada kebutuhan pengguna. Pengalaman tersebut tidak hanya meningkatkan kemampuan teknis dalam pengelolaan dan pengembangan sistem, tetapi juga mengembangkan kemampuan komunikasi, kerja sama tim, disiplin, serta profesionalisme kerja. Dengan demikian, kegiatan Praktik Kerja Lapangan di CV Newus Teknologi diharapkan dapat menjadi sarana untuk memperkecil kesenjangan antara teori yang diperoleh di perguruan tinggi dengan praktik yang diterapkan dalam industri teknologi informasi.")

    # 1.2
    add_subbab(doc, "1.2", "Tujuan PKL")
    add_body(doc, "Adapun tujuan dari pelaksanaan Praktik Kerja Lapangan (PKL) di CV Newus Teknologi adalah sebagai berikut:")

    for i, t in enumerate([
        "Memperoleh wawasan dan pengalaman kerja secara langsung dalam lingkungan industri teknologi informasi, khususnya pada proses pengelolaan dan pengembangan sistem informasi berbasis web.",
        "Menerapkan ilmu pengetahuan dan keterampilan yang telah diperoleh selama perkuliahan, terutama pada bidang pengembangan perangkat lunak, analisis kebutuhan sistem, serta pengelolaan basis data dalam lingkungan kerja yang nyata.",
        "Merancang dan membangun aplikasi Learning Management System (LMS) berbasis web yang mengintegrasikan frontend Next.js dengan backend RESTful API Laravel 11 dan basis data MySQL.",
        "Meningkatkan kemampuan komunikasi, kerja sama tim, manajemen waktu, dan pemecahan masalah yang diperlukan dalam lingkungan kerja profesional.",
        "Menambah wawasan mengenai proses bisnis perusahaan dan budaya kerja industri sebagai bekal dalam mempersiapkan diri menghadapi dunia kerja setelah lulus.",
    ], 1):
        add_numbered_item(doc, i, t)

    # 1.3
    add_subbab(doc, "1.3", "Kegunaan PKL")
    add_body(doc, "Pelaksanaan Praktik Kerja Lapangan (PKL) memberikan berbagai manfaat bagi mahasiswa, Fakultas Teknik dan Ilmu Komputer (FTIK) Universitas Teknokrat Indonesia, serta CV Newus Teknologi sebagai tempat pelaksanaan PKL. Adapun manfaat yang diperoleh adalah sebagai berikut:")

    add_body(doc, "1. Manfaat bagi Mahasiswa:")
    for i, m in enumerate([
        "Memperoleh pengalaman kerja secara langsung dalam lingkungan industri teknologi informasi profesional di CV Newus Teknologi.",
        "Mampu menerapkan ilmu pengetahuan dan keterampilan yang telah diperoleh selama perkuliahan ke dalam dunia kerja nyata, terutama pada bidang pengembangan perangkat lunak full-stack modern.",
        "Meningkatkan kemampuan teknis dalam perancangan arsitektur sistem, pemodelan basis data relasional, implementasi RESTful API, serta pengujian sistem otomatis.",
        "Mengembangkan kemampuan berpikir kritis dan pemecahan masalah dalam menghadapi berbagai kendala selama proses pengembangan dan implementasi sistem.",
        "Meningkatkan kemampuan komunikasi, kerja sama tim, kedisiplinan, serta manajemen waktu yang diperlukan dalam lingkungan kerja profesional.",
    ], 1):
        add_numbered_item(doc, i, m, indent=0.63)

    add_body(doc, "2. Manfaat bagi FTIK Universitas Teknokrat Indonesia:")
    for i, m in enumerate([
        "Menjadi sarana evaluasi terhadap kesesuaian kurikulum dan materi perkuliahan dengan kebutuhan nyata dunia industri teknologi informasi.",
        "Memperkuat jalinan kerja sama kemitraan antara FTIK Universitas Teknokrat Indonesia dengan CV Newus Teknologi sebagai industri mitra.",
        "Meningkatkan kualitas lulusan yang memiliki pengalaman kerja riil, kompetensi teknis teruji, serta kesiapan mental menghadapi dunia kerja.",
    ], 1):
        add_numbered_item(doc, i, m, indent=0.63)

    add_body(doc, "3. Manfaat bagi CV Newus Teknologi:")
    for i, m in enumerate([
        "Memperoleh kontribusi tenaga kerja mahasiswa dalam pengembangan produk Learning Management System sebagai solusi digital bagi institusi pendidikan mitra perusahaan.",
        "Mendukung percepatan penyelesaian target proyek operasional perusahaan melalui kontribusi nyata mahasiswa selama masa PKL.",
        "Menjadi sarana bagi perusahaan untuk mengenal, membina, dan mengevaluasi potensi mahasiswa sebagai calon talenta digital andal di masa mendatang.",
    ], 1):
        add_numbered_item(doc, i, m, indent=0.63)

    # 1.4
    add_subbab(doc, "1.4", "Tempat PKL")
    p = add_body_mixed(doc)
    ar(p, "Pelaksanaan Kegiatan Praktik Kerja Lapangan (PKL) dilaksanakan di CV Newus Teknologi yang berlokasi di Jl. Salim Batubara No.118, Kupang Teba, Kec. Teluk Betung Utara, Kota Bandar Lampung, Lampung 35212 (Website resmi: ", size=12)
    ar(p, "https://newus.id", italic=True, size=12)
    ar(p, "). Selama pelaksanaan PKL, penulis bertugas sebagai ", size=12)
    ar(p, "Software Engineer ", italic=True, size=12)
    ar(p, "yang bertanggung jawab dalam merancang dan mengembangkan ", size=12)
    ar(p, "Learning Management System ", italic=True, size=12)
    ar(p, "(LMS) menggunakan ", size=12)
    ar(p, "framework ", italic=True, size=12)
    ar(p, "Next.js dan Laravel berdasarkan kebutuhan yang ditentukan oleh perusahaan. Lokasi CV Newus Teknologi ditunjukkan pada Gambar 1.1.", size=12)

    add_gambar(doc, "Peta Lokasi CV Newus Teknologi dari Google Maps", "Gambar 1.1 Tempat Praktik Kerja Lapangan (PKL)")

    # 1.5
    add_subbab(doc, "1.5", "Jadwal Pelaksanaan PKL")
    add_body(doc, "Praktik Kerja Lapangan (PKL) dilaksanakan selama satu bulan terhitung mulai tanggal 6 Agustus 2026 sampai dengan 6 September 2026. Pembagian waktu pelaksanaan kegiatan telah disepakati antara pembimbing lapangan serta penulis dan disusun secara sistematis guna memastikan fokus serta optimalisasi penyelesaian tugas dapat tercapai secara maksimal. Jadwal dan tahapan kegiatan pelaksanaan Praktik Kerja Lapangan (PKL) dapat dilihat pada Tabel 1.1.")

    add_tabel_caption(doc, "Tabel 1.1 Tahapan Kegiatan Praktik Kerja Lapangan (PKL)")

    tbl_jadwal = doc.add_table(rows=6, cols=2)
    tbl_jadwal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl_jadwal.rows[0].cells
    for cell in hdr:
        set_cell_shading(cell, "D9E2F3")
    hp0 = hdr[0].paragraphs[0]
    hp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(hp0, "Tahapan Kegiatan", bold=True, size=11)
    hp1 = hdr[1].paragraphs[0]
    hp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(hp1, "Keterangan", bold=True, size=11)

    for i, (tahap, ket) in enumerate([
        ("Minggu ke-1\n(6-12 Agustus 2026)", "Orientasi perusahaan bersama Founder & CEO Ir. Antoni, S.T., IPM, penentuan topik proyek PKL, wawancara kebutuhan sistem, penyusunan Product Requirements Document (PRD), dan perancangan wireframe UI/UX di Figma."),
        ("Minggu ke-2\n(13-19 Agustus 2026)", "Perancangan skema database MySQL 11 tabel relasional, setup repository Git, inisialisasi framework Laravel 11 backend dan Next.js frontend, implementasi autentikasi Laravel Sanctum."),
        ("Minggu ke-3\n(20-26 Agustus 2026)", "Pengembangan modul kelas, materi ajar, penugasan LKPD, presensi mandiri berbasis jendela waktu, integrasi BroadcastChannel real-time sync, dan pembuatan antarmuka 3 dashboard peran pengguna."),
        ("Minggu ke-4\n(27 Agustus - 2 Sept 2026)", "Pengembangan modul rekapitulasi nilai rapor, input UTS/UAS, ekspor Excel, bulk import spreadsheet 50+ user, dan deployment produksi ke cloud Vercel dan Railway."),
        ("Minggu ke-5\n(3-6 September 2026)", "Eksekusi automated bot testing 53 skenario (100% passed), evaluasi hasil bersama tim pembimbing lapangan Newus Technology, penyusunan dokumentasi sistem, dan penulisan laporan akhir PKL."),
    ]):
        row = tbl_jadwal.rows[i+1]
        c0 = row.cells[0].paragraphs[0]
        c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(c0, tahap, size=10)
        c1 = row.cells[1].paragraphs[0]
        ar(c1, ket, size=10)

    # =========================================================================
    # BAB II TINJAUAN UMUM TEMPAT PKL
    # 2.1 Sejarah Perusahaan
    # 2.2 Struktur Organisasi
    # 2.3 Kegiatan Umum Perusahaan
    # =========================================================================
    doc.add_page_break()
    add_bab_title(doc, "II", "TINJAUAN UMUM TEMPAT PKL")

    add_subbab(doc, "2.1", "Sejarah Perusahaan")

    p = add_body_mixed(doc)
    ar(p, "CV Newus Teknologi (dikenal secara luas sebagai ", size=12)
    ar(p, "Newus Technology", bold=True, size=12)
    ar(p, ") merupakan perusahaan yang bergerak di bidang teknologi informasi dengan spesialisasi sebagai pengembang perangkat lunak (", size=12)
    ar(p, "Software Developer", italic=True, size=12)
    ar(p, ") dan konsultan teknologi informasi (", size=12)
    ar(p, "IT Consultant", italic=True, size=12)
    ar(p, ") profesional. Perusahaan ini didirikan pada tahun 2021 oleh ", size=12)
    ar(p, "Ir. Antoni, S.T., IPM ", bold=True, size=12)
    ar(p, "yang menjabat sebagai ", size=12)
    ar(p, "Founder & Chief Executive Officer (CEO). ", italic=True, size=12)
    ar(p, "Kantor pusat perusahaan berkedudukan di Jl. Salim Batubara No.118, Kupang Teba, Kec. Teluk Betung Utara, Kota Bandar Lampung, Lampung 35212 dengan portal informasi resmi dapat diakses melalui laman ", size=12)
    ar(p, "https://newus.id.", italic=True, size=12)

    p = add_body_mixed(doc)
    ar(p, "Sejak awal pendiriannya pada tahun 2021, Newus Technology memiliki komitmen kuat untuk memimpin pengembangan ekosistem dan sistem digital yang terukur (", size=12)
    ar(p, "scalable digital systems", italic=True, size=12)
    ar(p, ") bagi sektor instansi pemerintahan, fasilitas kesehatan (rumah sakit dan klinik), institusi pendidikan, korporasi swasta, serta pelaku Usaha Mikro, Kecil, dan Menengah (UMKM) di seluruh wilayah Indonesia. Hingga saat ini, Newus Technology telah berhasil menyelesaikan dan mengimplementasikan lebih dari 100 solusi teknologi digital, dengan fokus utama pada efisiensi operasional, transformasi digital terpadu, dan otomatisasi berbasis teknologi modern.", size=12)

    add_body(doc, "Visi utama CV Newus Teknologi adalah menjadi perusahaan konsultan dan pengembang perangkat lunak terdepan di Indonesia yang berdaya saing global serta mampu menghadirkan ekosistem teknologi praktis yang memberikan dampak terukur (measurable impact) dan keberlanjutan jangka panjang bagi kemajuan masyarakat, pemerintahan, dan industri.")

    add_body(doc, "Misi perusahaan diwujudkan melalui:")
    for i, m in enumerate([
        "Mengembangkan solusi perangkat lunak berkualitas tinggi dengan mengadopsi standar arsitektur sistem modern, performa tinggi, aman, dan mudah dikembangkan (scalable).",
        "Memberikan layanan konsultasi teknologi informasi yang solutif, transparan, dan memberikan nilai tambah nyata bagi transformasi digital para mitra bisnis dan instansi pemerintah.",
        "Membangun ekosistem kerja yang profesional, kolaboratif, dan inovatif yang mendukung pembinaan dan akselerasi talenta digital muda Indonesia.",
        "Mendukung digitalisasi sektor pendidikan melalui penyediaan platform e-learning, smart school management systems, dan integrasi data akademik terpusat.",
    ], 1):
        add_numbered_item(doc, i, m, indent=0.63)

    add_body(doc, "Logo resmi CV Newus Teknologi ditunjukkan pada Gambar 2.1.")
    add_gambar(doc, "Logo Resmi CV Newus Teknologi", "Gambar 2.1 Logo CV Newus Teknologi")

    # 2.2
    add_subbab(doc, "2.2", "Struktur Organisasi")
    p = add_body_mixed(doc)
    ar(p, "Struktur organisasi CV Newus Teknologi disusun secara profesional dan terencana guna mendukung kelancaran operasional, pembagian tugas yang jelas, dan pengawasan mutu pada setiap proyek teknologi yang dikerjakan. Perusahaan dipimpin oleh seorang ", size=12)
    ar(p, "Founder & CEO ", italic=True, size=12)
    ar(p, "yang membawahi divisi manajerial dan divisi teknis. Struktur organisasi perusahaan terdiri dari unit-unit kerja sebagai berikut:", size=12)

    add_numbered_item(doc, "1", "Direksi Utama (Executive Leadership): Dipimpin langsung oleh Ir. Antoni, S.T., IPM selaku Founder & CEO yang bertanggung jawab atas arah strategis perusahaan, kemitraan strategis, serta visi pengembangan produk teknologi.", indent=0.63)
    add_numbered_item(doc, "2", "Divisi Manajemen Proyek (Project Management): Dikoordinasikan oleh Project Manager (Retno Wardani dan Lulu Agustin) yang bertugas mengelola siklus proyek, memantau timeline, mengoordinasikan kebutuhan klien dengan tim teknis, serta memastikan kepatuhan terhadap Product Requirements Document (PRD).", indent=0.63)
    add_numbered_item(doc, "3", "Divisi Pengembangan Perangkat Lunak (Software Engineering & IT):", indent=0.63)
    add_body(doc, "a) Tim Frontend Developer: Dipimpin oleh Hamdan Habibi (Leader Front End Developer) dan Aldi Nugraha (Leader Front End), bersama Robi Hardinata, Bendry Lakburlawal, dan Ariston Rais Zidane, yang berfokus membangun antarmuka web modern, interaktif, responsif, dan performa tinggi menggunakan Next.js, React, dan Tailwind CSS.", first_indent=False)
    add_body(doc, "b) Tim Backend Developer & Database: Dipimpin oleh Muhammad Qomarudin (Leader Back-End Developer) dan Syukron Ma'mun (Tenaga Ahli Basis Data & Database), didukung Geraldine Firdaus, M. G Arma Yoga Pratama, dan Muhammad Muttaqin, yang bertanggung jawab atas perancangan database relasional, pengembangan RESTful API, keamanan autentikasi, dan skalabilitas server menggunakan Laravel, Node.js, dan MySQL.", first_indent=False)
    add_body(doc, "c) Tim UI/UX Design: Dipimpin oleh Qurrota Aini Dila (Leader UI/UX Designer) bersama Muhammad Raihan Puteranda dan Irsyad Abi Izzulhaq, yang bertugas melakukan riset pengguna, perancangan wireframe, user flow, dan high-fidelity prototype interaktif di Figma.", first_indent=False)
    add_body(doc, "d) Tim DevOps & Quality Assurance: Dikelola oleh Muhammad Nur Ashiddiqi (DevOps Engineer) yang bertanggung jawab atas infrastruktur cloud, deployment, dan otomatisasi CI/CD, serta Aditya Pangestu (Quality Assurance) yang bertugas menyusun skenario pengujian dan memastikan kualitas perangkat lunak bebas dari kecacatan (bug-free).", first_indent=False)
    add_numbered_item(doc, "4", "Divisi Operasional, Keuangan & Human Resources: Dikelola oleh Icha Aquinalda (Operasional Manager), Elvi Kholifatul Jannah (Admin Finance), dan Irene Nadia Vaniarinanta (Human Resources Officer) yang mengurus administrasi, keuangan, dan tata kelola SDM.", indent=0.63)
    add_numbered_item(doc, "5", "Divisi Business Development & Marketing: Dikelola oleh Yulivia Annisa Putri (Business Strategy Manager), Fadilla Pusvitasari (Business Development), dan Miacika Putri Paramitha (SEO Content Writer) yang bertanggung jawab atas perluasan pasar, strategi SEO, dan kemitraan klien.", indent=0.63)

    add_body(doc, "Bagan Struktur Organisasi CV Newus Teknologi ditunjukkan pada Gambar 2.2, dan daftar susunan tim inti perusahaan dirangkum pada Tabel 2.1.")
    add_gambar(doc, "Bagan Struktur Organisasi CV Newus Teknologi", "Gambar 2.2 Struktur Organisasi CV Newus Teknologi")

    add_tabel_caption(doc, "Tabel 2.1 Daftar Susunan Tim Inti CV Newus Teknologi")
    tbl_team = doc.add_table(rows=11, cols=4)
    tbl_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl_team.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
    for j, h in enumerate(["No", "Nama", "Jabatan / Posisi", "Divisi"]):
        cp = tbl_team.rows[0].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(cp, h, bold=True, size=10)

    team_summary = [
        ("1", "Ir. Antoni, S.T., IPM", "Founder & CEO", "Direksi Utama"),
        ("2", "Retno Wardani", "Project Manager", "Project Management"),
        ("3", "Hamdan Habibi", "Leader Front End Developer", "Frontend Development"),
        ("4", "Muhammad Qomarudin", "Leader Back-End Developer", "Backend Development"),
        ("5", "Syukron Ma'mun", "Tenaga Ahli Basis Data", "Database & Backend"),
        ("6", "Qurrota Aini Dila", "Leader UI/UX Designer", "UI/UX Design"),
        ("7", "Muhammad Nur Ashiddiqi", "DevOps Engineer", "Cloud & Infrastructure"),
        ("8", "Aditya Pangestu", "Quality Assurance Engineer", "QA & Testing"),
        ("9", "Irene Nadia Vaniarinanta", "Human Resources Officer", "Human Resources"),
        ("10", "Yulivia Annisa Putri", "Business Strategy Manager", "Business Development"),
    ]
    for i, (no, nama, pos, div) in enumerate(team_summary):
        row = tbl_team.rows[i+1]
        for j, val in enumerate([no, nama, pos, div]):
            cp = row.cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = ar(cp, val, size=10)
            if j == 1 and i == 0:
                r.bold = True

    # 2.3
    add_subbab(doc, "2.3", "Kegiatan Umum Perusahaan")
    add_body(doc, "Kegiatan umum CV Newus Teknologi berfokus pada penyediaan layanan teknologi informasi komprehensif dari hulu ke hilir, mulai dari konsultasi analisis kebutuhan, perancangan desain sistem, pengembangan kode program, pengujian mutu, hingga pemeliharaan dan deployment infrastruktur cloud. Portofolio layanan utama perusahaan meliputi:")

    kegiatan_list = [
        "Website & Web Application Development: Pembangunan aplikasi web interaktif skala korporat dan pemerintahan berbasis framework modern (Next.js, React, Laravel, Node.js). Contoh proyek nyata: SIINTAN Dinas Perkim Lampung Utara dan portal korporat terintegrasi.",
        "Mobile Application Development: Pengembangan aplikasi seluler native dan cross-platform (Flutter, React Native) untuk kebutuhan operasional lapangan, pemantauan kesehatan, dan layanan publik.",
        "Learning Management System (LMS) & Smart School Systems: Pembangunan platform pendidikan digital terpadu untuk sekolah dan perguruan tinggi yang mencakup manajemen kelas, e-learning, presensi online, penugasan digital, dan rekapitulasi nilai rapor otomatis.",
        "Custom Enterprise Software & ERP: Pengembangan perangkat lunak kustom sesuai kebutuhan unik alur bisnis klien, meliputi modul inventaris, logistik pergudangan, keuangan, dan manajemen aset.",
        "Cloud Infrastructure & DevOps Solutions: Perancangan arsitektur server cloud, deployment otomatis (CI/CD), konfigurasi load balancer, containerization (Docker), dan pemeliharaan server pada platform AWS, Google Cloud, Railway, dan Vercel.",
        "UI/UX Design & Human-Centered Research: Layanan riset pengalaman pengguna, perancangan user journey, pembuatan wireframe, design system, dan prototype interaktif berstandar industri.",
        "IT Consultant & Digital Transformation Advisory: Konsultasi strategi digital, audit sistem informasi, optimasi basis data, dan pendampingan transformasi digital bagi instansi pemerintah dan perusahaan swasta.",
    ]
    for i, k in enumerate(kegiatan_list, 1):
        add_numbered_item(doc, i, k)

    # =========================================================================
    # BAB III PELAKSANAAN PRAKTIK KERJA LAPANGAN
    # =========================================================================
    doc.add_page_break()
    add_bab_title(doc, "III", "PELAKSANAAN PRAKTIK KERJA LAPANGAN")

    p = add_body_mixed(doc)
    ar(p, "Pada bab ini dipaparkan pelaksanaan kerja masing-masing anggota kelompok PKL dalam mengembangkan ", size=12)
    ar(p, "Learning Management System ", italic=True, size=12)
    ar(p, "(LMS) di CV Newus Teknologi. Sistem yang dikembangkan menggunakan arsitektur ", size=12)
    ar(p, "decoupled ", italic=True, size=12)
    ar(p, "yang memisahkan sisi ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "(Next.js 14 dengan Tailwind CSS) dan ", size=12)
    ar(p, "backend ", italic=True, size=12)
    ar(p, "(Laravel 11 dengan MySQL 8.0). Arsitektur sistem secara menyeluruh ditunjukkan pada Gambar 3.1.", size=12)

    add_gambar(doc, "Arsitektur Sistem LMS menampilkan pemisahan Frontend (Next.js/Vercel) dan Backend (Laravel/Railway) yang berkomunikasi melalui RESTful API", "Gambar 3.1 Arsitektur Sistem LMS (Frontend dan Backend)")

    # ================================================================
    # KERTAS PENYEKAT BIRU - FERY DWI RAMADHI
    # ================================================================
    add_blue_separator(doc, "FERY DWI RAMADHI", "NPM 23312086")

    # 3.1 FERY DWI RAMADHI
    add_subbab(doc, "3.1", "FERY DWI RAMADHI (23312086)")

    add_subbab(doc, "3.1.1", "Bidang Kerja")
    p = add_body_mixed(doc)
    ar(p, "Selama melaksanakan Praktik Kerja Lapangan di CV Newus Teknologi, penulis dipercaya mengemban posisi sebagai ", size=12)
    ar(p, "Full-Stack Developer ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "Project Lead. ", italic=True, size=12)
    ar(p, "Tanggung jawab utama penulis meliputi perancangan arsitektur sistem menyeluruh, perancangan dan implementasi basis data relasional MySQL, pembangunan ", size=12)
    ar(p, "RESTful API ", italic=True, size=12)
    ar(p, "menggunakan ", size=12)
    ar(p, "framework ", italic=True, size=12)
    ar(p, "Laravel 11, integrasi keamanan autentikasi menggunakan Laravel Sanctum, serta penerapan mekanisme sinkronisasi data ", size=12)
    ar(p, "real-time ", italic=True, size=12)
    ar(p, "antar-tab browser. Penulis juga berperan sebagai koordinator tim yang bertanggung jawab membagi tugas, melakukan ", size=12)
    ar(p, "code review, ", italic=True, size=12)
    ar(p, "dan memastikan integrasi seluruh modul berjalan dengan baik.", size=12)

    add_subbab(doc, "3.1.2", "Pelaksanaan Kerja")
    add_body(doc, "Pelaksanaan kerja yang dilakukan oleh penulis selama masa PKL dijabarkan ke dalam beberapa tahapan sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Pada tahap pertama, penulis merancang basis data relasional yang terdiri atas 11 tabel terintegrasi: ", size=12)
    ar(p, "users, courses, course_student, materials, assignments, submissions, attendances, notifications, settings, activity_logs, ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "personal_access_tokens. ", italic=True, size=12)
    ar(p, "Setiap entitas tabel dikonstruksi menggunakan ", size=12)
    ar(p, "Laravel Eloquent Migration ", italic=True, size=12)
    ar(p, "dengan penegakan integritas data melalui ", size=12)
    ar(p, "Foreign Key Constraints ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "Cascade Deletion. ", italic=True, size=12)
    ar(p, "Perancangan ini dilakukan dengan menyesuaikan kebutuhan fungsional yang telah diidentifikasi pada tahap analisis kebutuhan sistem. Dokumentasi diagram basis data ditunjukkan pada Gambar 3.2.", size=12)

    add_gambar(doc, "Entity Relationship Diagram (ERD) Basis Data LMS dengan 11 tabel relasional", "Gambar 3.2 Entity Relationship Diagram (ERD) Basis Data")

    add_tabel_caption(doc, "Tabel 3.2 Struktur Tabel Basis Data MySQL")
    tbl_db = doc.add_table(rows=12, cols=3)
    tbl_db.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl_db.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
    for j, h in enumerate(["No", "Nama Tabel", "Keterangan"]):
        cp = tbl_db.rows[0].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(cp, h, bold=True, size=10)
    
    db_tables = [
        ("1", "users", "Data pengguna (admin, guru, siswa) dengan role dan autentikasi"),
        ("2", "courses", "Data kelas/mata pelajaran dengan relasi ke guru pengampu"),
        ("3", "course_student", "Tabel pivot relasi many-to-many kelas dan siswa"),
        ("4", "materials", "Modul bahan ajar (judul, konten, file lampiran)"),
        ("5", "assignments", "Penugasan LKPD dengan deadline dan bobot nilai"),
        ("6", "submissions", "Pengumpulan tugas siswa dengan file dan nilai"),
        ("7", "attendances", "Data presensi siswa per kelas per tanggal"),
        ("8", "notifications", "Notifikasi real-time untuk seluruh pengguna"),
        ("9", "settings", "Konfigurasi sistem (nama sekolah, logo, dll)"),
        ("10", "activity_logs", "Catatan aktivitas pengguna untuk audit trail"),
        ("11", "personal_access_tokens", "Token autentikasi Laravel Sanctum"),
    ]
    for i, (no, nama, ket) in enumerate(db_tables):
        row = tbl_db.rows[i+1]
        for j, val in enumerate([no, nama, ket]):
            cp = row.cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = ar(cp, val, size=10)
            if j == 1:
                r.italic = True

    p = add_body_mixed(doc)
    ar(p, "Selanjutnya, penulis membangun lebih dari 50 ", size=12)
    ar(p, "endpoint REST API ", italic=True, size=12)
    ar(p, "yang terorganisir di bawah ", size=12)
    ar(p, "namespace ", italic=True, size=12)
    ar(p, "/api/v1. Seluruh rute dilindungi oleh ", size=12)
    ar(p, "middleware auth:sanctum ", italic=True, size=12)
    ar(p, "dengan otorisasi berbasis peran pengguna (", size=12)
    ar(p, "Role-Based Access Control / RBAC", italic=True, size=12)
    ar(p, ") yang membedakan hak akses Administrator, Guru Pengajar, dan Peserta Didik. Mekanisme ", size=12)
    ar(p, "token Bearer ", italic=True, size=12)
    ar(p, "digunakan untuk menjaga keamanan sesi tanpa membebani server. Daftar endpoint utama ditunjukkan pada Tabel 3.1.", size=12)

    add_tabel_caption(doc, "Tabel 3.1 Daftar Endpoint RESTful API Laravel 11")
    tbl_api = doc.add_table(rows=11, cols=4)
    tbl_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl_api.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
    for j, h in enumerate(["No", "Method", "Endpoint", "Keterangan"]):
        cp = tbl_api.rows[0].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(cp, h, bold=True, size=9)
    
    api_data = [
        ("1", "POST", "/api/v1/login", "Autentikasi pengguna, generate token"),
        ("2", "GET", "/api/v1/courses", "Daftar kelas sesuai peran pengguna"),
        ("3", "POST", "/api/v1/courses", "Buat kelas baru (Admin/Guru)"),
        ("4", "GET", "/api/v1/materials", "Daftar modul bahan ajar per kelas"),
        ("5", "POST", "/api/v1/assignments", "Buat penugasan LKPD (Guru)"),
        ("6", "POST", "/api/v1/submissions", "Submit jawaban tugas (Siswa)"),
        ("7", "POST", "/api/v1/attendances", "Catat presensi mandiri (Siswa)"),
        ("8", "GET", "/api/v1/notifications", "Ambil notifikasi pengguna"),
        ("9", "GET", "/api/v1/reports/grades", "Rekapitulasi nilai rapor"),
        ("10", "POST", "/api/v1/users/import", "Bulk import akun dari Excel"),
    ]
    for i, (no, method, ep, ket) in enumerate(api_data):
        row = tbl_api.rows[i+1]
        for j, val in enumerate([no, method, ep, ket]):
            cp = row.cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT
            r = ar(cp, val, size=9)
            if j == 2:
                r.italic = True

    p = add_body_mixed(doc)
    ar(p, "Penulis juga merancang arsitektur sinkronisasi mutasi data pada ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "Next.js. Setiap kali terjadi operasi ", size=12)
    ar(p, "Create, Update, ", italic=True, size=12)
    ar(p, "atau ", size=12)
    ar(p, "Delete, ", italic=True, size=12)
    ar(p, "API ", size=12)
    ar(p, "helper ", italic=True, size=12)
    ar(p, "akan memicu ", size=12)
    ar(p, "event BroadcastChannel ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "Custom DOM Event. Custom hook useRealtimeData ", italic=True, size=12)
    ar(p, "kemudian menangkap sinyal tersebut dan melakukan ", size=12)
    ar(p, "revalidation ", italic=True, size=12)
    ar(p, "data secara instan tanpa memerlukan ", size=12)
    ar(p, "reload ", italic=True, size=12)
    ar(p, "halaman web secara manual. Mekanisme ini memastikan bahwa setiap perubahan data yang dilakukan oleh satu pengguna pada satu tab browser akan langsung tercermin di seluruh tab lain yang sedang membuka halaman yang sama.", size=12)

    add_body(doc, "Tampilan halaman login sistem LMS yang telah dikembangkan ditunjukkan pada Gambar 3.5. Halaman login dirancang dengan antarmuka yang bersih dan responsif, dilengkapi dengan fitur demo akun yang memungkinkan pengguna mencoba sistem tanpa harus mendaftar terlebih dahulu.")
    add_gambar(doc, "Tampilan Halaman Login Sistem LMS dengan fitur Demo Accounts", "Gambar 3.5 Tampilan Halaman Login Sistem LMS")

    p = add_body_mixed(doc)
    ar(p, "Penulis melakukan ", size=12)
    ar(p, "deployment backend ", italic=True, size=12)
    ar(p, "Laravel ke ", size=12)
    ar(p, "cloud platform ", italic=True, size=12)
    ar(p, "Railway menggunakan ", size=12)
    ar(p, "runtime ", italic=True, size=12)
    ar(p, "PHP 8.2 dan MySQL 8.0 terkelola, serta menghubungkannya dengan domain Vercel ", size=12)
    ar(p, "frontend. ", italic=True, size=12)
    ar(p, "Konfigurasi ", size=12)
    ar(p, "environment variables ", italic=True, size=12)
    ar(p, "dilakukan dengan cermat untuk memastikan koneksi antara ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "backend ", italic=True, size=12)
    ar(p, "berjalan lancar pada lingkungan produksi. Pengujian fungsional seluruh ", size=12)
    ar(p, "endpoint ", italic=True, size=12)
    ar(p, "dilakukan secara berkala menggunakan Postman Collection untuk memastikan keandalan API.", size=12)

    # 3.1.3
    add_subbab(doc, "3.1.3", "Kendala yang Dihadapi")
    add_body(doc, "Selama pelaksanaan Praktik Kerja Lapangan (PKL) di CV Newus Teknologi, penulis menghadapi beberapa kendala yang menjadi bagian dari proses adaptasi terhadap lingkungan kerja profesional serta teknologi yang digunakan. Adapun kendala yang dihadapi adalah sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Kendala pertama yang dihadapi penulis berkaitan dengan inkonsistensi penulisan huruf besar/kecil (", size=12)
    ar(p, "case-sensitivity", italic=True, size=12)
    ar(p, ") pada nilai status presensi siswa saat validasi ", size=12)
    ar(p, "controller backend. ", italic=True, size=12)
    ar(p, "Karena nilai yang dikirim dari ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "terkadang menggunakan huruf kapital sedangkan validasi ", size=12)
    ar(p, "enum ", italic=True, size=12)
    ar(p, "di Laravel bersifat ", size=12)
    ar(p, "case-sensitive, ", italic=True, size=12)
    ar(p, "hal ini menyebabkan kesalahan validasi yang tidak terduga pada saat pengguna melakukan presensi.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Kendala kedua yang dihadapi penulis berkaitan dengan konfigurasi ", size=12)
    ar(p, "Cross-Origin Resource Sharing ", italic=True, size=12)
    ar(p, "(CORS) dan penanganan ", size=12)
    ar(p, "session cookies ", italic=True, size=12)
    ar(p, "antar-domain saat ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "Vercel berkomunikasi dengan ", size=12)
    ar(p, "backend ", italic=True, size=12)
    ar(p, "Railway. Permintaan API dari domain Vercel ke domain Railway awalnya ditolak oleh ", size=12)
    ar(p, "browser ", italic=True, size=12)
    ar(p, "karena kebijakan ", size=12)
    ar(p, "Same-Origin Policy, ", italic=True, size=12)
    ar(p, "sehingga proses autentikasi dan pertukaran data tidak dapat dilakukan.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Kendala ketiga yang dihadapi penulis berkaitan dengan kebutuhan pembaruan data tampilan di sisi klien yang sebelumnya mengharuskan pengguna melakukan ", size=12)
    ar(p, "refresh browser ", italic=True, size=12)
    ar(p, "secara berulang untuk melihat perubahan terbaru. Hal ini mengurangi kenyamanan penggunaan sistem dan tidak sesuai dengan standar pengalaman pengguna (", size=12)
    ar(p, "User Experience", italic=True, size=12)
    ar(p, ") aplikasi web modern.", size=12)

    # 3.1.4
    add_subbab(doc, "3.1.4", "Cara Mengatasi Kendala")
    add_body(doc, "Dalam menghadapi berbagai kendala selama pelaksanaan PKL, penulis berupaya untuk terus meningkatkan pemahaman dan kemampuan teknis melalui diskusi dengan pembimbing lapangan serta studi mandiri. Adapun cara yang dilakukan untuk mengatasi setiap kendala adalah sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala pertama terkait inkonsistensi penulisan status presensi, penulis menambahkan normalisasi ", size=12)
    ar(p, "string ", italic=True, size=12)
    ar(p, "menggunakan fungsi ", size=12)
    ar(p, "strtolower(trim($request->status)) ", italic=True, size=12)
    ar(p, "pada ", size=12)
    ar(p, "AttendanceController ", italic=True, size=12)
    ar(p, "Laravel sebelum validasi ", size=12)
    ar(p, "enum ", italic=True, size=12)
    ar(p, "dijalankan. Dengan demikian, sistem dapat menerima berbagai variasi penulisan huruf besar maupun kecil dengan aman tanpa menimbulkan kesalahan validasi.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala kedua terkait CORS, penulis mengonfigurasi file ", size=12)
    ar(p, "config/cors.php ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "config/sanctum.php ", italic=True, size=12)
    ar(p, "dengan mendaftarkan ", size=12)
    ar(p, "origin ", italic=True, size=12)
    ar(p, "domain ", size=12)
    ar(p, "frontend ", italic=True, size=12)
    ar(p, "serta menyematkan ", size=12)
    ar(p, "header Accept: application/json ", italic=True, size=12)
    ar(p, "pada ", size=12)
    ar(p, "client ", italic=True, size=12)
    ar(p, "Axios/Fetch di sisi ", size=12)
    ar(p, "frontend. ", italic=True, size=12)
    ar(p, "Selain itu, penulis juga menambahkan konfigurasi ", size=12)
    ar(p, "supports_credentials ", italic=True, size=12)
    ar(p, "pada Laravel agar ", size=12)
    ar(p, "cookie ", italic=True, size=12)
    ar(p, "autentikasi dapat dikirim lintas-domain.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala ketiga terkait pembaruan data tanpa ", size=12)
    ar(p, "reload, ", italic=True, size=12)
    ar(p, "penulis membangun ", size=12)
    ar(p, "custom hook useRealtimeData ", italic=True, size=12)
    ar(p, "berbasis ", size=12)
    ar(p, "BroadcastChannel ", italic=True, size=12)
    ar(p, "browser API yang secara otomatis mendeteksi perubahan data dan memicu ", size=12)
    ar(p, "query revalidation ", italic=True, size=12)
    ar(p, "saat tab kembali aktif. Mekanisme ini memungkinkan pembaruan data secara instan tanpa campur tangan pengguna, sehingga pengalaman penggunaan sistem menjadi lebih responsif dan modern.", size=12)

    # ================================================================
    # KERTAS PENYEKAT BIRU - FATHUR RAMANTHA
    # ================================================================
    add_blue_separator(doc, "FATHUR RAMANTHA", "NPM 23312105")

    # 3.2 FATHUR RAMANTHA
    add_subbab(doc, "3.2", "FATHUR RAMANTHA (23312105)")
    add_subbab(doc, "3.2.1", "Bidang Kerja")

    p = add_body_mixed(doc)
    ar(p, "Dalam pelaksanaan PKL, penulis bertindak sebagai ", size=12)
    ar(p, "Frontend Developer ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "UI/UX Specialist. ", italic=True, size=12)
    ar(p, "Fokus tugas penulis adalah menyusun rancangan antarmuka pengguna berstandar modern, mengonversi desain Figma ke dalam kode komponen Next.js (App Router), menerapkan sistem tata warna dan responsivitas berbasis Tailwind CSS, serta mengoptimalkan pengalaman interaksi pengguna di semua perangkat. Penulis juga bertanggung jawab memastikan seluruh halaman yang dibangun memiliki konsistensi visual dan dapat diakses dengan baik pada berbagai ukuran layar.", size=12)

    add_subbab(doc, "3.2.2", "Pelaksanaan Kerja")
    p = add_body_mixed(doc)
    ar(p, "Tahapan pelaksanaan kerja yang diselesaikan oleh penulis meliputi perancangan ", size=12)
    ar(p, "wireframe ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "high-fidelity UI ", italic=True, size=12)
    ar(p, "di Figma untuk seluruh halaman sistem, implementasi komponen Next.js dan Tailwind CSS, serta pembuatan antarmuka interaktif untuk tiga peran pengguna. Proses perancangan dimulai dengan pembuatan ", size=12)
    ar(p, "wireframe low-fidelity ", italic=True, size=12)
    ar(p, "untuk menentukan tata letak dan alur navigasi, kemudian dilanjutkan dengan pembuatan ", size=12)
    ar(p, "high-fidelity mockup ", italic=True, size=12)
    ar(p, "yang menyertakan palet warna, tipografi, dan komponen visual akhir. Tampilan ", size=12)
    ar(p, "dashboard ", italic=True, size=12)
    ar(p, "Administrator ditunjukkan pada Gambar 3.6.", size=12)

    add_gambar(doc, "Tampilan Dashboard Administrator dengan statistik dan grafik", "Gambar 3.6 Tampilan Dashboard Administrator")

    p = add_body_mixed(doc)
    ar(p, "Penulis membangun lebih dari 30 halaman dan 20 komponen ", size=12)
    ar(p, "reusable ", italic=True, size=12)
    ar(p, "Next.js yang dirancang untuk dapat digunakan kembali di berbagai bagian sistem. Komponen-komponen tersebut mencakup ", size=12)
    ar(p, "Sidebar, Navbar ", italic=True, size=12)
    ar(p, "dengan lonceng notifikasi ", size=12)
    ar(p, "real-time, StatCard, CourseCard, ", italic=True, size=12)
    ar(p, "modal dialog interaktif, dan tabel data responsif. Setiap komponen dikembangkan dengan pendekatan ", size=12)
    ar(p, "atomic design ", italic=True, size=12)
    ar(p, "untuk memastikan modularitas dan kemudahan pemeliharaan kode. Daftar komponen utama ditunjukkan pada Tabel 3.3.", size=12)

    add_tabel_caption(doc, "Tabel 3.3 Daftar Komponen Antarmuka Next.js")
    tbl_comp = doc.add_table(rows=11, cols=3)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl_comp.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
    for j, h in enumerate(["No", "Nama Komponen", "Fungsi"]):
        cp = tbl_comp.rows[0].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(cp, h, bold=True, size=10)
    
    comps = [
        ("1", "Sidebar", "Navigasi utama dashboard sesuai peran"),
        ("2", "Navbar", "Header dengan notifikasi dan profil pengguna"),
        ("3", "StatCard", "Kartu statistik ringkasan dashboard"),
        ("4", "CourseCard", "Kartu tampilan kelas dengan info ringkas"),
        ("5", "DataTable", "Tabel data responsif dengan pagination"),
        ("6", "ModalDialog", "Dialog interaktif untuk CRUD operations"),
        ("7", "FileUploader", "Komponen unggah file drag-and-drop"),
        ("8", "SkeletonLoader", "Placeholder loading state komponen"),
        ("9", "NotificationBell", "Lonceng notifikasi real-time"),
        ("10", "GradeInputForm", "Formulir input nilai UTS/UAS"),
    ]
    for i, (no, nama, fungsi) in enumerate(comps):
        row = tbl_comp.rows[i+1]
        for j, val in enumerate([no, nama, fungsi]):
            cp = row.cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = ar(cp, val, size=10)
            if j == 1:
                r.italic = True

    add_body(doc, "Tampilan dashboard Guru Pengajar yang menampilkan informasi kelas yang diampu, statistik kehadiran, dan akses cepat ke modul penugasan ditunjukkan pada Gambar 3.7.")
    add_gambar(doc, "Tampilan Dashboard Guru Pengajar dengan daftar kelas dan statistik", "Gambar 3.7 Tampilan Dashboard Guru Pengajar")

    add_body(doc, "Penulis mengembangkan halaman katalog kelas yang memungkinkan siswa melihat daftar kelas yang telah didaftarkan, mengakses materi pembelajaran, serta mengikuti kegiatan presensi dan penugasan. Tampilan dashboard Siswa ditunjukkan pada Gambar 3.8.")
    add_gambar(doc, "Tampilan Dashboard Siswa dengan katalog kelas dan jadwal tugas", "Gambar 3.8 Tampilan Dashboard Siswa")

    p = add_body_mixed(doc)
    ar(p, "Selain itu, penulis juga membangun halaman manajemen kelas yang memungkinkan administrator dan guru untuk mengelola data kelas, menambahkan siswa ke kelas, serta mengatur pengaturan kelas. Halaman ini dilengkapi dengan fitur pencarian dan filter untuk memudahkan pengelolaan data. Tampilan halaman manajemen kelas ditunjukkan pada Gambar 3.9.", size=12)
    add_gambar(doc, "Tampilan Halaman Manajemen Kelas dengan daftar kelas dan aksi", "Gambar 3.9 Tampilan Halaman Manajemen Kelas")

    add_body(doc, "Penulis mengembangkan halaman presensi siswa yang menampilkan riwayat kehadiran dan tombol presensi mandiri satu klik. Sistem presensi dirancang dengan mekanisme jendela waktu yang hanya mengizinkan siswa melakukan presensi dalam rentang waktu yang telah ditentukan oleh guru pengajar. Tampilan halaman presensi ditunjukkan pada Gambar 3.10.")
    add_gambar(doc, "Tampilan Halaman Presensi Siswa dengan tombol presensi mandiri", "Gambar 3.10 Tampilan Halaman Presensi Siswa")

    add_body(doc, "Penulis juga membangun halaman penugasan LKPD yang memungkinkan guru membuat tugas dengan lampiran file, menetapkan deadline, dan memberikan penilaian. Siswa dapat mengunggah jawaban melalui antarmuka drag-and-drop yang intuitif. Tampilan halaman penugasan ditunjukkan pada Gambar 3.11.")
    add_gambar(doc, "Tampilan Halaman Penugasan LKPD dengan form submit jawaban", "Gambar 3.11 Tampilan Halaman Penugasan LKPD")

    # 3.2.3
    add_subbab(doc, "3.2.3", "Kendala Yang Dihadapi")
    add_body(doc, "Selama pelaksanaan PKL, penulis menghadapi beberapa kendala teknis yang berkaitan dengan pengembangan antarmuka pengguna. Adapun kendala yang dihadapi adalah sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Kendala pertama yang dihadapi penulis adalah timbulnya pergeseran tata letak antarmuka (", size=12)
    ar(p, "Cumulative Layout Shift / CLS", italic=True, size=12)
    ar(p, ") saat data statistik grafik pertama kali dimuat dari server. Kondisi ini menyebabkan elemen-elemen pada halaman bergeser secara tiba-tiba sehingga mengganggu kenyamanan pengguna dalam berinteraksi dengan sistem.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Kendala kedua yang dihadapi penulis berkaitan dengan tampilan tabel rekapitulasi nilai dan presensi yang terpotong pada layar beresolusi kecil (ponsel dan tablet). Hal ini menyebabkan pengguna tidak dapat melihat seluruh data yang ditampilkan tanpa harus melakukan ", size=12)
    ar(p, "scroll ", italic=True, size=12)
    ar(p, "secara horizontal, yang tidak selalu intuitif bagi pengguna.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Kendala ketiga yang dihadapi penulis berkaitan dengan kompleksitas penanganan status tombol (", size=12)
    ar(p, "disabled state ", italic=True, size=12)
    ar(p, "dan animasi ", size=12)
    ar(p, "loading", italic=True, size=12)
    ar(p, ") saat pengiriman formulir berlangsung. Tanpa penanganan yang tepat, pengguna dapat menekan tombol ", size=12)
    ar(p, "submit ", italic=True, size=12)
    ar(p, "berkali-kali dan menyebabkan duplikasi data.", size=12)

    # 3.2.4
    add_subbab(doc, "3.2.4", "Cara Mengatasi Kendala")
    add_body(doc, "Dalam menghadapi kendala-kendala tersebut, penulis melakukan berbagai upaya perbaikan sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala pertama, penulis menerapkan ", size=12)
    ar(p, "Skeleton Loader Components ", italic=True, size=12)
    ar(p, "dengan dimensi pasti sebelum data selesai di-", size=12)
    ar(p, "fetch ", italic=True, size=12)
    ar(p, "dari API, sehingga antarmuka tetap stabil dan tidak terjadi pergeseran tata letak yang mengganggu.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala kedua, penulis menerapkan pembungkus ", size=12)
    ar(p, "horizontal overflow-x-auto ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "styling container queries ", italic=True, size=12)
    ar(p, "pada tabel laporan agar data tetap dapat di-", size=12)
    ar(p, "scroll ", italic=True, size=12)
    ar(p, "dengan mulus di perangkat ", size=12)
    ar(p, "mobile ", italic=True, size=12)
    ar(p, "tanpa memotong konten yang ditampilkan.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala ketiga, penulis mengintegrasikan ", size=12)
    ar(p, "state isSubmitting ", italic=True, size=12)
    ar(p, "dengan animasi ", size=12)
    ar(p, "spinner ", italic=True, size=12)
    ar(p, "Lucide-React pada seluruh tombol aksi. Ketika formulir sedang diproses, tombol secara otomatis dinonaktifkan untuk mencegah pengiriman ganda.", size=12)

    # ================================================================
    # KERTAS PENYEKAT BIRU - I PUTU PANDU WIRANATA
    # ================================================================
    add_blue_separator(doc, "I PUTU PANDU WIRANATA", "NPM 23312088")

    # 3.3 I PUTU PANDU WIRANATA
    add_subbab(doc, "3.3", "I PUTU PANDU WIRANATA (23312088)")
    add_subbab(doc, "3.3.1", "Bidang Kerja")

    p = add_body_mixed(doc)
    ar(p, "Dalam kegiatan PKL, penulis bertugas sebagai ", size=12)
    ar(p, "System Analyst, Database Administrator ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "Quality Assurance (QA). ", italic=True, size=12)
    ar(p, "Ruang lingkup kerja penulis mencakup analisis kebutuhan fungsional sistem, penyusunan dokumentasi teknis berupa diagram UML, manajemen ", size=12)
    ar(p, "seeder ", italic=True, size=12)
    ar(p, "data massal untuk pengujian, pembuatan fitur ", size=12)
    ar(p, "import spreadsheet, ", italic=True, size=12)
    ar(p, "serta perancangan dan eksekusi pengujian otomatis menyeluruh terhadap seluruh fitur sistem.", size=12)

    add_subbab(doc, "3.3.2", "Pelaksanaan Kerja")
    p = add_body_mixed(doc)
    ar(p, "Pada tahap awal, penulis menyusun ", size=12)
    ar(p, "Product Requirements Document ", italic=True, size=12)
    ar(p, "(PRD) yang mendokumentasikan keseluruhan kebutuhan fungsional dan non-fungsional sistem LMS. Dokumen ini mencakup spesifikasi fitur, prioritas pengembangan, serta kriteria penerimaan (", size=12)
    ar(p, "acceptance criteria", italic=True, size=12)
    ar(p, ") untuk setiap modul. Selain itu, penulis mendesain Diagram ", size=12)
    ar(p, "Use Case ", italic=True, size=12)
    ar(p, "terpadu yang mencakup 50 fitur fungsional untuk tiga peran pengguna (Administrator, Guru Pengajar, dan Peserta Didik). Diagram ", size=12)
    ar(p, "Use Case ", italic=True, size=12)
    ar(p, "ditunjukkan pada Gambar 3.3.", size=12)

    add_gambar(doc, "Diagram Use Case Terpadu Sistem LMS dengan 50 fitur fungsional", "Gambar 3.3 Diagram Use Case Terpadu Sistem LMS")

    p = add_body_mixed(doc)
    ar(p, "Selanjutnya, penulis menyusun Diagram Alir Sistem (", size=12)
    ar(p, "Flowchart", italic=True, size=12)
    ar(p, ") terpadu yang menggambarkan alur kerja keseluruhan sistem mulai dari proses ", size=12)
    ar(p, "login, ", italic=True, size=12)
    ar(p, "pengarahan berdasarkan peran, hingga akses ke modul-modul fungsional. Diagram Alir ini menjadi acuan bagi tim pengembang dalam memastikan konsistensi logika bisnis di seluruh bagian sistem. Diagram Alir Sistem ditunjukkan pada Gambar 3.4.", size=12)

    add_gambar(doc, "Diagram Alir (Flowchart) Sistem Terpadu menampilkan alur login dan navigasi berdasarkan peran", "Gambar 3.4 Diagram Alir (Flowchart) Sistem Terpadu")

    p = add_body_mixed(doc)
    ar(p, "Penulis menyusun skrip ", size=12)
    ar(p, "Database Seeder ", italic=True, size=12)
    ar(p, "untuk menginisialisasi 61 akun pengguna nyata yang terdiri dari 1 akun administrator, 13 akun guru pengampu mata pelajaran, dan 47 akun siswa. Selain itu, penulis juga membuat data ", size=12)
    ar(p, "seed ", italic=True, size=12)
    ar(p, "untuk 6 kelas aktif dan 3 tugas LKPD sebagai data awal pengujian. Penulis turut merancang modul ", size=12)
    ar(p, "import ", italic=True, size=12)
    ar(p, "akun massal dari file Excel (.xlsx/.csv) yang memungkinkan administrator mendaftarkan puluhan hingga ratusan akun pengguna dalam satu operasi. Tampilan halaman ", size=12)
    ar(p, "bulk import ", italic=True, size=12)
    ar(p, "ditunjukkan pada Gambar 3.17.", size=12)

    add_gambar(doc, "Tampilan Halaman Bulk Import Spreadsheet untuk pendaftaran akun massal", "Gambar 3.17 Tampilan Halaman Bulk Import Spreadsheet")

    add_body(doc, "Penulis juga terlibat dalam pengembangan modul rekapitulasi nilai rapor yang mengintegrasikan data dari penugasan LKPD, ujian UTS, dan ujian UAS. Modul ini dilengkapi dengan fitur ekspor ke format Excel (.xlsx) untuk memudahkan pelaporan akademik. Tampilan halaman rekapitulasi nilai ditunjukkan pada Gambar 3.12.")
    add_gambar(doc, "Tampilan Halaman Rekapitulasi Nilai Rapor dengan kolom LKPD, UTS, UAS", "Gambar 3.12 Tampilan Halaman Rekapitulasi Nilai Rapor")

    add_body(doc, "Tampilan lonceng notifikasi yang memberikan pemberitahuan secara real-time kepada pengguna ditunjukkan pada Gambar 3.13.")
    add_gambar(doc, "Tampilan Lonceng Notifikasi Real-time pada Navbar", "Gambar 3.13 Tampilan Lonceng Notifikasi Real-time")

    add_body(doc, "Tampilan fitur ekspor laporan ke format Excel yang memungkinkan guru dan administrator mengunduh data nilai dan kehadiran ditunjukkan pada Gambar 3.14.")
    add_gambar(doc, "Tampilan Fitur Ekspor Laporan Excel (.xlsx)", "Gambar 3.14 Tampilan Fitur Ekspor Laporan Excel")

    add_body(doc, "Penulis juga mengembangkan halaman manajemen modul bahan ajar yang memungkinkan guru mengunggah materi pembelajaran berupa teks, file PDF, dan dokumen pendukung lainnya. Tampilan halaman manajemen modul bahan ajar ditunjukkan pada Gambar 3.15.")
    add_gambar(doc, "Tampilan Halaman Manajemen Modul Bahan Ajar", "Gambar 3.15 Tampilan Halaman Manajemen Modul Bahan Ajar")

    add_body(doc, "Tampilan halaman input nilai UTS/UAS yang memungkinkan guru memberikan penilaian ujian semester ditunjukkan pada Gambar 3.16.")
    add_gambar(doc, "Tampilan Halaman Input Nilai UTS/UAS oleh Guru", "Gambar 3.16 Tampilan Halaman Input Nilai UTS/UAS")

    p = add_body_mixed(doc)
    ar(p, "Pada tahap akhir, penulis merancang dan mengeksekusi skrip ", size=12)
    ar(p, "bot ", italic=True, size=12)
    ar(p, "penguji otomatis menggunakan Node.js yang mensimulasikan interaksi nyata 4 akun pengguna (Admin, Guru, Siswa 1, Siswa 2) terhadap 8 kategori fitur inti (53 skenario uji coba) secara menyeluruh terhadap server ", size=12)
    ar(p, "backend ", italic=True, size=12)
    ar(p, "Railway. Seluruh skenario pengujian berhasil dieksekusi dengan tingkat keberhasilan 100%. Hasil pengujian otomatis ditunjukkan pada Tabel 3.4.", size=12)

    add_tabel_caption(doc, "Tabel 3.4 Hasil Pengujian Otomatis (E2E Automated Bot Testing)")
    tbl_test = doc.add_table(rows=9, cols=4)
    tbl_test.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl_test.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
    for j, h in enumerate(["No", "Modul yang Diuji", "Akun Penguji", "Hasil"]):
        cp = tbl_test.rows[0].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(cp, h, bold=True, size=10)

    for i, (no, modul, akun, hasil) in enumerate([
        ("1", "Autentikasi dan Token Sanctum", "Admin, Guru, Siswa", "Passed"),
        ("2", "Dashboard Statistik Global", "Administrator", "Passed"),
        ("3", "Manajemen Akun Pengguna", "Administrator", "Passed"),
        ("4", "Bulk Import Spreadsheet", "Administrator", "Passed"),
        ("5", "Manajemen Kelas dan Materi", "Guru Pengajar", "Passed"),
        ("6", "Presensi Mandiri Siswa", "Siswa Pembelajar", "Passed"),
        ("7", "Penugasan LKPD dan Penilaian", "Guru, Siswa", "Passed"),
        ("8", "Rekapitulasi Nilai Rapor", "Admin, Guru, Siswa", "Passed"),
    ]):
        row = tbl_test.rows[i+1]
        for j, val in enumerate([no, modul, akun, hasil]):
            cp = row.cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
            r = ar(cp, val, size=10)
            if j == 3:
                r.bold = True

    add_body(doc, "Dokumentasi kegiatan Praktik Kerja Lapangan (PKL) di CV Newus Teknologi ditunjukkan pada Gambar 3.18.")
    add_gambar(doc, "Dokumentasi Kegiatan PKL di CV Newus Teknologi (foto bersama tim)", "Gambar 3.18 Dokumentasi Kegiatan PKL di CV Newus Teknologi")

    # 3.3.3
    add_subbab(doc, "3.3.3", "Kendala Yang Dihadapi")
    add_body(doc, "Selama pelaksanaan PKL, penulis menghadapi beberapa kendala yang berkaitan dengan proses analisis data, manajemen basis data, dan pengujian sistem. Adapun kendala yang dihadapi adalah sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Kendala pertama yang dihadapi penulis adalah format file Excel yang diunggah pengguna seringkali memiliki urutan kolom atau nama ", size=12)
    ar(p, "header ", italic=True, size=12)
    ar(p, "yang tidak seragam. Variasi format ini menyebabkan proses ", size=12)
    ar(p, "parsing ", italic=True, size=12)
    ar(p, "data pada ", size=12)
    ar(p, "controller import ", italic=True, size=12)
    ar(p, "gagal karena sistem tidak dapat mengenali kolom yang dimaksud.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Kendala kedua yang dihadapi penulis berkaitan dengan skenario presensi siswa yang gagal ketika diuji di luar rentang jam operasional kelas yang ditentukan oleh guru. Hal ini menunjukkan bahwa logika validasi waktu pada ", size=12)
    ar(p, "backend ", italic=True, size=12)
    ar(p, "perlu disempurnakan agar memberikan pesan kesalahan yang informatif kepada pengguna.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Kendala ketiga yang dihadapi penulis berkaitan dengan waktu pengujian manual fitur yang sangat lama jika harus ", size=12)
    ar(p, "login ", italic=True, size=12)
    ar(p, "dan ", size=12)
    ar(p, "logout ", italic=True, size=12)
    ar(p, "bergantian dengan 4 akun berbeda untuk menguji skenario lintas-peran. Proses ini memakan waktu yang tidak efisien dan rentan terhadap kesalahan manusia (", size=12)
    ar(p, "human error", italic=True, size=12)
    ar(p, ").", size=12)

    # 3.3.4
    add_subbab(doc, "3.3.4", "Cara Mengatasi Kendala")
    add_body(doc, "Dalam menghadapi kendala-kendala tersebut, penulis melakukan berbagai upaya perbaikan sebagai berikut:")

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala pertama, penulis menyediakan ", size=12)
    ar(p, "template ", italic=True, size=12)
    ar(p, "file master dan menambahkan validasi ", size=12)
    ar(p, "header ", italic=True, size=12)
    ar(p, "otomatis pada ", size=12)
    ar(p, "controller import ", italic=True, size=12)
    ar(p, "Laravel menggunakan pustaka Maatwebsite Excel. Sistem akan menampilkan pesan kesalahan yang spesifik apabila format kolom tidak sesuai dengan ", size=12)
    ar(p, "template ", italic=True, size=12)
    ar(p, "yang telah ditetapkan.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala kedua, penulis menambahkan logika validasi zona waktu server menggunakan pustaka Carbon (Asia/Jakarta) dengan pesan ", size=12)
    ar(p, "error ", italic=True, size=12)
    ar(p, "edukatif jika presensi diakses di luar jadwal. Pesan kesalahan dirancang agar memberikan informasi yang jelas mengenai rentang waktu presensi yang diperbolehkan.", size=12)

    p = add_body_mixed(doc)
    ar(p, "Untuk mengatasi kendala ketiga, penulis mengembangkan skrip ", size=12)
    ar(p, "bot ", italic=True, size=12)
    ar(p, "pengujian otomatis menggunakan Node.js yang menjalankan 53 ", size=12)
    ar(p, "test case ", italic=True, size=12)
    ar(p, "secara otomatis dalam waktu kurang dari 15 detik dengan hasil 100% ", size=12)
    ar(p, "Passed. ", italic=True, size=12)
    ar(p, "Skrip ini mensimulasikan seluruh alur penggunaan sistem secara menyeluruh tanpa memerlukan intervensi manual.", size=12)

    # =========================================================================
    # BAB IV PENUTUP
    # =========================================================================
    doc.add_page_break()
    add_bab_title(doc, "IV", "PENUTUP")

    add_subbab(doc, "4.1", "Simpulan")
    add_body(doc, "Berdasarkan hasil pelaksanaan Praktik Kerja Lapangan (PKL) yang dilaksanakan di CV Newus Teknologi selama satu bulan (6 Agustus sampai dengan 6 September 2026), dapat disimpulkan beberapa hal sebagai berikut:")

    for i, k in enumerate([
        "Pelaksanaan Praktik Kerja Lapangan memberikan pengalaman kerja secara langsung kepada penulis dalam lingkungan industri teknologi informasi profesional di CV Newus Teknologi, khususnya dalam pengembangan Learning Management System berbasis web menggunakan framework Next.js dan Laravel 11.",
        "Penulis mampu menerapkan ilmu yang diperoleh selama perkuliahan pada proses pengembangan perangkat lunak secara profesional, mulai dari analisis kebutuhan sistem, perancangan basis data relasional 11 tabel, implementasi lebih dari 50 endpoint RESTful API, pembangunan antarmuka pengguna responsif untuk tiga peran pengguna, hingga pengujian dan deployment sistem ke lingkungan cloud produksi.",
        "Sistem Learning Management System yang dikembangkan telah diuji secara komprehensif menggunakan Automated End-to-End Bot Testing dengan 53 skenario uji coba lintas-peran dan berhasil meraih tingkat keberhasilan 100% (Passed), yang menunjukkan bahwa seluruh fitur sistem telah berfungsi sesuai dengan spesifikasi kebutuhan yang ditentukan.",
        "Aplikasi telah berhasil di-deploy ke lingkungan cloud produksi publik (Frontend di Vercel dan Backend di Railway) dan siap diimplementasikan oleh institusi pendidikan mitra CV Newus Teknologi sebagai platform pembelajaran daring yang terintegrasi.",
        "Selama pelaksanaan PKL, penulis berhasil meningkatkan kemampuan komunikasi, kerja sama tim, manajemen waktu, berpikir logis, pemecahan masalah, kedisiplinan, serta tanggung jawab dalam menyelesaikan pekerjaan secara profesional.",
    ], 1):
        add_numbered_item(doc, i, k)

    add_subbab(doc, "4.2", "Saran")
    add_body(doc, "Berdasarkan hasil pelaksanaan Praktik Kerja Lapangan (PKL), penulis memberikan beberapa saran sebagai berikut:")

    add_body(doc, "1. Untuk CV Newus Teknologi:")
    for i, s in enumerate([
        "Perusahaan dapat terus mempertahankan sistem pembelajaran, mentoring, dan pendampingan yang telah berjalan dengan sangat baik sehingga peserta PKL selanjutnya dapat memahami alur kerja dan standar industri secara optimal.",
        "Perusahaan dapat menambahkan modul video conference interaktif dan payment gateway otomatis ke dalam platform LMS untuk memperluas portofolio produk e-learning dan meningkatkan daya saing pasar.",
    ], 1):
        add_numbered_item(doc, i, s, indent=0.63)

    add_body(doc, "2. Kontribusi terhadap IPTEK:")
    for i, s in enumerate([
        "Penggunaan arsitektur decoupled (Next.js dan Laravel) dalam pengembangan LMS dapat dijadikan referensi bagi penelitian dan pengembangan sistem informasi pendidikan berbasis web di masa mendatang.",
        "Mekanisme sinkronisasi data real-time menggunakan BroadcastChannel API dapat diadopsi dan dikembangkan lebih lanjut untuk aplikasi web lain yang membutuhkan pembaruan data tanpa reload halaman.",
        "Universitas Teknokrat Indonesia dapat memperbanyak materi praktikum berbasis arsitektur web modern dan cloud deployment pada kurikulum perkuliahan agar mahasiswa memiliki kesiapan yang lebih baik dalam menghadapi tuntutan industri.",
    ], 1):
        add_numbered_item(doc, i, s, indent=0.63)

    # =========================================================================
    # DAFTAR PUSTAKA (Lampiran 7 - Sistem Harvard, spasi 1)
    # =========================================================================
    doc.add_page_break()
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12, keep_with_next=True)
    ar(p, "DAFTAR PUSTAKA", bold=True, size=14)

    pustaka_list = [
        'Brauer, F., Castillo-Chavez, C., 2001, Mathematical Models in Population Biology and Epidemiology, Springer-Verlag, Inc., New York.',
        'Fathansyah, 2018, Basis Data (Edisi Revisi), Informatika Bandung, Bandung.',
        'FTIK-UTI, 2018, Buku Panduan Praktik Kerja Lapangan (PKL), TCTC Universitas Teknokrat Indonesia, Bandar Lampung.',
        'Laravel LLC, 2024, Laravel Documentation (Version 11.x), https://laravel.com/docs, diakses 15 Agustus 2026.',
        'Newus Technology, 2024, Tim & Profil Perusahaan Newus Technology, https://newus.id/team, diakses 24 Agustus 2026.',
        'Next.js Team, 2024, Next.js App Router Documentation, Vercel Inc., https://nextjs.org/docs, diakses 15 Agustus 2026.',
        'Pressman, R. S. dan Maxim, B. R., 2020, Software Engineering: A Practitioner\'s Approach (9th ed.), McGraw-Hill Education, New York.',
        'Rosa, A. S. dan Shalahuddin, M., 2019, Rekayasa Perangkat Lunak Terstruktur dan Berorientasi Objek, Informatika Bandung, Bandung.',
        'Schwaber, K. dan Sutherland, J., 2020, The Scrum Guide: The Definitive Guide to Scrum: The Rules of the Game, Scrum.org.',
        'Tailwind Labs, 2024, Tailwind CSS Documentation (Version 3.x), https://tailwindcss.com/docs, diakses 15 Agustus 2026.',
        'Turban, E., et al., 2005, Decision Support System and Intelligent System, Edisi 7 jilid 1, Pearson Education Inc., Upper Saddle River, New Jersey dan Penerbit Andi, Yogyakarta.',
    ]
    for pust in pustaka_list:
        p = doc.add_paragraph()
        sfmt(p, line_spacing=1.0, space_after=6, left_indent=1.27, first_line_indent=-1.27)
        ar(p, pust, size=12)

    # =========================================================================
    # LAMPIRAN
    # =========================================================================
    doc.add_page_break()
    p = doc.add_paragraph()
    sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12, keep_with_next=True)
    ar(p, "LAMPIRAN", bold=True, size=14)

    for lamp_num, lamp_judul, lamp_placeholder in [
        ("Lampiran 1", "Formulir Penilaian Fery Dwi Ramadhi", "Formulir Penilaian Mahasiswa PKL - Fery Dwi Ramadhi"),
        ("Lampiran 2", "Catatan Harian Fery Dwi Ramadhi", "Catatan Harian PKL - Fery Dwi Ramadhi"),
        ("Lampiran 3", "Formulir Penilaian Fathur Ramantha", "Formulir Penilaian Mahasiswa PKL - Fathur Ramantha"),
        ("Lampiran 4", "Catatan Harian Fathur Ramantha", "Catatan Harian PKL - Fathur Ramantha"),
        ("Lampiran 5", "Formulir Penilaian I Putu Pandu Wiranata", "Formulir Penilaian Mahasiswa PKL - I Putu Pandu Wiranata"),
        ("Lampiran 6", "Catatan Harian I Putu Pandu Wiranata", "Catatan Harian PKL - I Putu Pandu Wiranata"),
    ]:
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=12)
        ar(p, f"{lamp_num}  {lamp_judul}", bold=True, size=12)
        p = doc.add_paragraph()
        sfmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_before=12)
        ar(p, f"[Gambar untuk {lamp_placeholder}]", italic=True, size=11)
        doc.add_page_break()

    # SAVE
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs", "Laporan_PKL_FTIK_Teknokrat_LMS_2026.docx")
    doc.save(output_path)
    print(f"SUCCESS: Laporan PKL dengan Data Resmi Newus Technology berhasil disimpan di: {output_path}")

if __name__ == "__main__":
    create_laporan()
